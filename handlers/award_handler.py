"""수상보고서 처리 핸들러"""

import asyncio
import os
import re
import tempfile
import time
from datetime import datetime

import pytz
import requests
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from google.oauth2.service_account import Credentials
from googleapiclient.discovery import build
from telegram import Update
from telegram.ext import ContextTypes

import config
from handlers.report_parser_utils import parse_multiline_kv, extract_first_line_meta

KST = pytz.timezone('Asia/Seoul')

AWARD_SPREADSHEET_ID = '1MM79Y5rjOT-s8GnN1WGfnRb3Bq5iZA-Ro4fQzEGZoB4'
AWARD_SHEET_NAME = '수상보고창'
AWARD_GROUP_ID = -1002777848839
AWARD_TOPIC_ID = 3553
AWARD_RECIPIENT_ID = 754270008

HEADERS = ['등록일시', '지역', '지부', '수상명', '수상일시', '장소',
           '수여자', '수상자', '수상내용',
           '사진1링크', '사진2링크', '사진3링크', '사진4링크', '사진5링크',
           '사진6링크', '사진7링크', '사진8링크', '사진9링크', '사진10링크']

AWARD_PENDING_REPORTS = {}   # (chat_id, topic_id, user_id) -> {data, photos, saved, last_photo_time, created, origin}
AWARD_PENDING_PHOTOS = {}    # (chat_id, topic_id, user_id) -> {photos, created}
AWARD_MEDIA_CACHE = {}       # media_group_id -> {photos, caption, processed, created, key, origin}
AWARD_PHOTOS_TTL = 600       # 10분 — 여러 앨범 누적 시간 고려
MAX_PHOTOS = 10

ADMIN_USER_ID = 97057565


def _alias_hint(field: str) -> str:
    aliases = AWARD_KEY_ALIASES.get(field, [field]) if 'AWARD_KEY_ALIASES' in globals() else [field]
    return f"{field} (다음 중 하나로 입력 가능: {', '.join(aliases)})"


async def _notify_admin(context, error_text: str):
    try:
        await context.bot.send_message(chat_id=ADMIN_USER_ID, text=error_text[:4000])
    except Exception as e:
        print(f"❌ 관리자 알림 실패: {e}")


async def _send_to_recipient(context, **kwargs):
    """서무 DM 전송, 실패 시 관리자 백업 알림"""
    try:
        if 'document' in kwargs:
            await context.bot.send_document(chat_id=AWARD_RECIPIENT_ID, **kwargs)
        else:
            await context.bot.send_message(chat_id=AWARD_RECIPIENT_ID, **kwargs)
        return True
    except Exception as e:
        await _notify_admin(context, f"❌ 수상보고서 서무 DM 실패: {e}\n내용: {str(kwargs.get('text', ''))[:300]}")
        return False


# ── 파싱 ─────────────────────────────────────────────────────────────────────

AWARD_KEY_ALIASES = {
    '수상명': ['수상명', '보고제목', '내용', '행사명', '표창명', '상명'],
    '수상일시': ['수상일시', '행사일시', '일시', '일자', '날짜'],
    '수여자': ['수여자', '시상자'],
    '수상자': ['수상자'],
    '수상내용': ['수상내용', '수상내역', '내역'],
    '장소': ['장소', '행사장소'],
}
REQUIRED_FIELDS = {'수상명', '수상일시', '수상자'}


def parse_award_caption(caption: str) -> dict | None:
    if not caption:
        return None
    if '수상' not in caption or '보고' not in caption:
        return None
    lines = [l.strip() for l in caption.strip().splitlines() if l.strip()]
    if not lines:
        return None
    지역, 지부 = extract_first_line_meta(
        lines[0], ['수상보고서', '수상보고', '보고서', '보고']
    )
    body = '\n'.join(lines[1:])
    meta = parse_multiline_kv(body, AWARD_KEY_ALIASES)
    data: dict = {
        '등록일시': datetime.now(KST).strftime('%Y-%m-%d %H:%M:%S'),
        '지역': 지역, '지부': 지부, '사진링크': '',
    }
    data.update(meta)
    missing = [f for f in REQUIRED_FIELDS if not data.get(f)]
    if missing:
        data['_missing'] = missing
    return data


# ── Google Sheets ─────────────────────────────────────────────────────────────

def _get_sheet_service():
    scopes = config.get('google_scopes', ['https://www.googleapis.com/auth/spreadsheets'])
    creds = Credentials.from_service_account_file('serviceAccountKey.json', scopes=scopes)
    return build('sheets', 'v4', credentials=creds)


def _ensure_header(service):
    result = service.spreadsheets().values().get(
        spreadsheetId=AWARD_SPREADSHEET_ID,
        range=f'{AWARD_SHEET_NAME}!A1:Z1'
    ).execute()
    existing = result.get('values', [[]])
    existing_row = existing[0] if existing else []
    if not existing_row or not any(existing_row):
        service.spreadsheets().values().update(
            spreadsheetId=AWARD_SPREADSHEET_ID,
            range=f'{AWARD_SHEET_NAME}!A1',
            valueInputOption='RAW',
            body={'values': [HEADERS]}
        ).execute()
        return
    if len(existing_row) < len(HEADERS):
        # 사진6~10링크 컬럼 추가 마이그레이션
        service.spreadsheets().values().update(
            spreadsheetId=AWARD_SPREADSHEET_ID,
            range=f'{AWARD_SHEET_NAME}!A1',
            valueInputOption='RAW',
            body={'values': [HEADERS]}
        ).execute()
        print(f"✅ 수상보고창 헤더 마이그레이션: {len(existing_row)} → {len(HEADERS)}컬럼")


def save_to_sheet(data: dict) -> bool:
    try:
        service = _get_sheet_service()
        _ensure_header(service)
        row = [data.get(h, '') for h in HEADERS]
        service.spreadsheets().values().append(
            spreadsheetId=AWARD_SPREADSHEET_ID,
            range=f'{AWARD_SHEET_NAME}!A:S',
            valueInputOption='RAW',
            insertDataOption='INSERT_ROWS',
            body={'values': [row]}
        ).execute()
        print(f"✅ 수상보고서 저장 완료: {data.get('수상명')}")
        return True
    except Exception as e:
        print(f"❌ 수상보고서 저장 오류: {e}")
        return False


# ── 사진 다운로드 ─────────────────────────────────────────────────────────────

def download_photo(url: str) -> str | None:
    try:
        token = config.get('telegram_token')
        full_url = url if url.startswith('http') else f"https://api.telegram.org/file/bot{token}/{url}"
        resp = requests.get(full_url, timeout=10)
        if resp.status_code == 200:
            suffix = '.jpg' if 'jpg' in url.lower() else '.png'
            tmp = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
            tmp.write(resp.content)
            tmp.close()
            return tmp.name
        return None
    except Exception as e:
        print(f"❌ 사진 다운로드 오류: {e}")
        return None


# ── Word 생성 ─────────────────────────────────────────────────────────────────

def _set_cell_bg(cell, color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)


def _set_para_bg(para, color: str):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    pPr.append(shd)


def _label(cell, text: str):
    cell.text = text
    run = cell.paragraphs[0].runs[0]
    run.bold = True
    run.font.size = Pt(10)
    _set_cell_bg(cell, 'D5E8F0')


def _value(cell, text: str):
    cell.text = text or '-'
    cell.paragraphs[0].runs[0].font.size = Pt(10)


def generate_docx(data: dict, photo_paths, output_path: str) -> bool:
    """photo_paths: list[str] (여러 장) 또는 단일 str/None (하위 호환)"""
    if photo_paths is None:
        photo_paths = []
    elif isinstance(photo_paths, str):
        photo_paths = [photo_paths]
    try:
        doc = Document()
        for section in doc.sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

        # 1. 제목
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_p.add_run(
            f"{data.get('지역', '')} {data.get('지부', '')}\n자원봉사 수상보고서"
        )
        title_run.bold = True
        title_run.font.size = Pt(16)
        title_run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        title_p.paragraph_format.space_after = Pt(12)

        # 2. 정보 테이블 (4열: label/value/label/value)
        table = doc.add_table(rows=0, cols=4)
        table.style = 'Table Grid'
        table.columns[0].width = Cm(3)
        table.columns[1].width = Cm(5.5)
        table.columns[2].width = Cm(3)
        table.columns[3].width = Cm(4.5)

        # 행 1: 수상명 (값 3열 병합)
        r0 = table.add_row()
        _label(r0.cells[0], '수상명')
        r0.cells[1].merge(r0.cells[3])
        _value(r0.cells[1], data.get('수상명'))

        # 행 2: 장소 | 수상일시
        r1 = table.add_row()
        _label(r1.cells[0], '장소')
        _value(r1.cells[1], data.get('장소'))
        _label(r1.cells[2], '수상일시')
        _value(r1.cells[3], data.get('수상일시'))

        # 행 3: 수여자 | 수상자
        r2 = table.add_row()
        _label(r2.cells[0], '수여자')
        _value(r2.cells[1], data.get('수여자'))
        _label(r2.cells[2], '수상자')
        _value(r2.cells[3], data.get('수상자'))

        # 행 4: 수상내용 (값 3열 병합)
        r3 = table.add_row()
        _label(r3.cells[0], '수상내용')
        r3.cells[1].merge(r3.cells[3])
        _value(r3.cells[1], data.get('수상내용'))

        doc.add_paragraph()

        # 3. 사진 섹션 (다중 사진 그리드)
        from handlers.report_base import add_photos_grid
        valid_paths = [p for p in photo_paths if p and os.path.exists(p)]
        if valid_paths:
            add_photos_grid(doc, valid_paths, title='상장, 상패 사진')
        else:
            photo_title = doc.add_paragraph()
            photo_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pt_run = photo_title.add_run('상장, 상패 사진')
            pt_run.bold = True
            pt_run.font.size = Pt(11)
            _set_para_bg(photo_title, 'CCCCCC')
            doc.add_paragraph('[사진 없음]').alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

        # 4. 하단
        footer_p = doc.add_paragraph()
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_p.add_run('위와 같이 보고합니다.\n').font.size = Pt(11)

        date_p = doc.add_paragraph()
        date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_p.add_run(data.get('수상일시', '')).font.size = Pt(11)

        reporter_p = doc.add_paragraph()
        reporter_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        reporter_run = reporter_p.add_run(
            f"{data.get('지역', '')} {data.get('지부', '')}"
        )
        reporter_run.font.size = Pt(11)

        doc.save(output_path)
        print(f"✅ 수상보고서 Word 생성 완료: {output_path}")
        return True

    except Exception as e:
        print(f"❌ Word 생성 오류: {e}")
        return False


# ── 사진 대기 관리 ─────────────────────────────────────────────────────────────

def _award_key(chat_id: int, thread_id, user_id: int = 0):
    """동시 제출 충돌 방지: chat+topic+user 조합"""
    return (chat_id, thread_id or 0, user_id or 0)


def _key_to_str(key) -> str:
    if len(key) == 3:
        return f"{key[0]}:{key[1]}:{key[2]}"
    return f"{key[0]}:{key[1]}:0"


def _save_pending_to_db(key, entry: dict):
    try:
        from database import save_pending_report
        save_pending_report(
            'award', _key_to_str(key),
            entry.get('data', {}),
            entry.get('photos', []),
            entry.get('last_photo_time', 0),
            entry.get('created', time.time()),
        )
    except Exception as e:
        print(f"⚠️ award PENDING DB 저장 실패: {e}")


def _delete_pending_from_db(key):
    try:
        from database import delete_pending_report
        delete_pending_report('award', _key_to_str(key))
    except Exception:
        pass


def _save_pending_photos_to_db(key, photos: list):
    try:
        from database import save_pending_photos_db
        save_pending_photos_db('award', _key_to_str(key), photos, time.time())
    except Exception:
        pass


def _parse_key_str(s: str):
    """저장 키 문자열을 튜플로. 구버전(chat:topic)도 호환."""
    parts = s.split(':')
    if len(parts) == 3:
        return (int(parts[0]), int(parts[1]), int(parts[2]))
    if len(parts) == 2:
        return (int(parts[0]), int(parts[1]), 0)
    raise ValueError(f"bad key: {s}")


def restore_pending_from_db():
    """봇 시작 시 호출 — DB에서 PENDING 상태 복원"""
    try:
        from database import load_pending_reports, load_pending_photos_db
        for r in load_pending_reports('award'):
            try:
                key = _parse_key_str(r['pending_key'])
                AWARD_PENDING_REPORTS[key] = {
                    'data': r['data'], 'photos': r['photos'], 'saved': False,
                    'last_photo_time': r['last_photo_time'], 'created': r['created'],
                    'origin': r['data'].get('_origin', {}),
                }
            except Exception:
                continue
        for r in load_pending_photos_db('award'):
            try:
                key = _parse_key_str(r['pending_key'])
                AWARD_PENDING_PHOTOS[key] = {
                    'photos': r['photos'], 'created': r['created'],
                }
            except Exception:
                continue
        if AWARD_PENDING_REPORTS or AWARD_PENDING_PHOTOS:
            print(f"✅ 수상보고서 PENDING 복원: 보고서 {len(AWARD_PENDING_REPORTS)}건, 사진 {len(AWARD_PENDING_PHOTOS)}건")
    except Exception as e:
        print(f"⚠️ 수상보고서 PENDING 복원 실패: {e}")


def _cleanup_award_photos():
    now = time.time()
    expired = [k for k, v in list(AWARD_PENDING_PHOTOS.items())
               if now - v.get('created', 0) > AWARD_PHOTOS_TTL]
    for k in expired:
        AWARD_PENDING_PHOTOS.pop(k, None)
        try:
            from database import delete_pending_photos_db
            delete_pending_photos_db('award', _key_to_str(k))
        except Exception:
            pass


async def _finalize_award(context, data: dict, photos: list, *, origin: dict = None):
    """트랜잭션: 사진 → Word → 시트 → DM 전송 → 보고자 reply.
    D-1 엄격: 사진 1장이라도 실패하면 처리 차단."""
    from handlers.report_base import (
        download_photos_batch, send_to_recipient as base_send,
        notify_admin as base_notify, with_sheet_retry, reply_to_origin,
    )
    from database import log_report_stage, record_submission
    if origin is None:
        origin = data.get('_origin', {}) or {}
    user_id = origin.get('user_id')
    output_path = None
    tmp_files = []
    try:
        loop = asyncio.get_running_loop()

        # ── 1. 사진 다운로드 (D-1 엄격) ────────────────────────────────
        tmp_files, photo_failed = await download_photos_batch(photos[:MAX_PHOTOS])
        log_report_stage(
            'award', 'photos_downloaded',
            'ok' if photo_failed == 0 else 'fail',
            user_id=user_id,
            detail=f"ok={len(tmp_files)} fail={photo_failed}"
        )
        if photo_failed > 0:
            for tmp in tmp_files:
                try: os.remove(tmp)
                except Exception: pass
            await reply_to_origin(
                context.bot, origin,
                f"❌ 사진 다운로드 실패\n"
                f"사진 {len(photos)}장 중 {photo_failed}장 실패\n"
                f"모든 사진을 다시 보내주세요"
            )
            log_report_stage('award', 'finalize', 'fail',
                             user_id=user_id, detail='photo_failed')
            return

        # ── 2. Word 생성 ────────────────────────────────────────────────
        now_str = datetime.now(KST).strftime('%Y%m%d_%H%M%S')
        output_path = f"/tmp/award_{now_str}.docx"
        try:
            docx_ok = await loop.run_in_executor(
                None, generate_docx, data, tmp_files, output_path
            )
        except Exception as e:
            docx_ok = False
            print(f"❌ 수상 Word 생성 예외: {e}")
        log_report_stage('award', 'docx_generated',
                         'ok' if docx_ok else 'fail', user_id=user_id)

        # 사진 임시파일 정리 (Word 생성 후)
        for tmp in tmp_files:
            try: os.remove(tmp)
            except Exception: pass

        if not docx_ok:
            try:
                if output_path and os.path.exists(output_path):
                    os.remove(output_path)
            except Exception: pass
            await reply_to_origin(
                context.bot, origin,
                "❌ Word 파일 생성 실패\n"
                "사진과 보고서를 다시 보내주세요\n"
                "시트 저장도 안 됐습니다"
            )
            log_report_stage('award', 'finalize', 'fail',
                             user_id=user_id, detail='docx_fail')
            return

        # ── 3. 시트 저장 (Word 검증 통과 후) ────────────────────────────
        for i in range(1, MAX_PHOTOS + 1):
            data[f'사진{i}링크'] = photos[i-1] if i <= len(photos) else ''
        try:
            sheet_ok = await loop.run_in_executor(
                None, with_sheet_retry, save_to_sheet, data, 3
            )
        except Exception as e:
            sheet_ok = False
            print(f"❌ 수상 시트 저장 예외: {e}")
        log_report_stage('award', 'sheet_saved',
                         'ok' if sheet_ok else 'fail',
                         user_id=user_id, chat_id=origin.get('chat_id'),
                         topic_id=origin.get('message_thread_id'),
                         message_id=origin.get('message_id'))

        # ── 4. 서무 DM 요약 ────────────────────────────────────────────
        if sheet_ok:
            summary = (
                f"✅ 수상보고서 처리 완료\n"
                f"📌 {data.get('지역')} {data.get('지부')}\n"
                f"🏆 {data.get('수상명')}\n"
                f"📅 {data.get('수상일시')}\n"
                f"📸 사진 {len(photos)}장 첨부"
            )
        else:
            summary = (
                f"⚠️ {data.get('지역')} {data.get('지부')} 수상보고서 - 시트 저장 실패!\n"
                f"Word 파일은 정상이지만 스프레드시트 자동 저장 실패\n"
                f"수동으로 수상보고창 시트에 추가 부탁드립니다\n\n"
                f"🏆 {data.get('수상명')}\n"
                f"📅 {data.get('수상일시')}\n"
                f"📸 사진 {len(photos)}장"
            )
        dm_ok = await base_send(context.bot, AWARD_RECIPIENT_ID, text=summary)
        log_report_stage('award', 'recipient_dm_sent',
                         'ok' if dm_ok else 'fail', user_id=user_id)

        # ── 5. Word 파일 서무에게 전송 ─────────────────────────────────
        if output_path and os.path.exists(output_path):
            filename = (
                f"{data.get('지역', '')}_{data.get('지부', '')}_수상보고서"
                f"_{datetime.now(KST).strftime('%Y%m%d')}.docx"
            ).replace(' ', '_').replace('/', '-')
            try:
                with open(output_path, 'rb') as f:
                    await base_send(
                        context.bot, AWARD_RECIPIENT_ID,
                        document=f, filename=filename,
                        caption="📄 수상보고서 Word 파일"
                    )
            finally:
                try: os.remove(output_path)
                except Exception: pass

        # 중복 기록
        try:
            sub_hash = f"{data.get('지부', '')}|{data.get('수상명', '')}"
            record_submission('award', sub_hash, summary[:200])
        except Exception: pass

        # ── 6. 보고자 최종 reply ────────────────────────────────────────
        if sheet_ok and dm_ok:
            await reply_to_origin(
                context.bot, origin,
                f"✅ 모든 처리 완료\n📸 사진 {len(photos)}장 첨부됨"
            )
        elif sheet_ok and not dm_ok:
            await reply_to_origin(
                context.bot, origin,
                "⚠️ 시트 저장 ✅\n⚠️ 서무 DM 전송 실패 - 관리자에게 알림됨"
            )
        else:
            await reply_to_origin(
                context.bot, origin,
                "⚠️ 시트 저장 실패 - 자동 처리됨\n"
                "Word 파일은 서무에게 정상 전송됐습니다\n"
                "시트는 서무가 수동 추가합니다"
            )
        log_report_stage('award', 'reporter_ack_sent', 'ok', user_id=user_id)
    except Exception as e:
        import traceback
        # 임시 파일 정리
        for tmp in tmp_files:
            try: os.remove(tmp)
            except Exception: pass
        if output_path:
            try: os.remove(output_path)
            except Exception: pass
        log_report_stage('award', 'finalize', 'fail',
                         user_id=user_id, detail=str(e)[:200])
        await base_notify(
            context.bot,
            f"❌ 수상보고서 처리 중 예외 발생: {e}\n"
            f"데이터: {str(data)[:300]}\n"
            f"{traceback.format_exc()[:1000]}"
        )
        await reply_to_origin(
            context.bot, origin,
            f"❌ 처리 중 오류 발생: {str(e)[:200]}\n관리자에게 알림됨"
        )


async def _flush_award_report(context, key):
    """60초 기다린 후, 사진이 계속 오면 5초씩 연장 (최대 5분)"""
    start = time.time()
    await asyncio.sleep(60)
    while True:
        entry = AWARD_PENDING_REPORTS.get(key)
        if not entry or entry.get('saved'):
            return
        last_photo = entry.get('last_photo_time', 0)
        elapsed = time.time() - start
        if last_photo > 0 and (time.time() - last_photo) < 5 and elapsed < 300:
            await asyncio.sleep(5)
            continue
        break
    entry = AWARD_PENDING_REPORTS.pop(key, None)
    if not entry or entry.get('saved'):
        return
    entry['saved'] = True
    _delete_pending_from_db(key)
    await _finalize_award(context, entry['data'], entry['photos'],
                          origin=entry.get('origin'))


# ── 메인 핸들러 ───────────────────────────────────────────────────────────────

def _award_parse_and_check(text: str, source_label: str, bot_coro_factory):
    """parse 후 누락 필드 확인. (코루틴 팩토리 패턴 대신 직접 반환)"""
    pass  # 아래 핸들러에서 인라인으로 처리


async def _process_award_album(context, media_group_id: str):
    """앨범(media_group) 사진 모두 도착 대기 후 finalize"""
    await asyncio.sleep(3)
    cache = AWARD_MEDIA_CACHE.get(media_group_id)
    if not cache or cache.get('processed'):
        return
    AWARD_MEDIA_CACHE[media_group_id]['processed'] = True
    caption = cache.get('caption', '')
    photos = cache.get('photos', [])[:MAX_PHOTOS]
    key = cache.get('key')
    origin = cache.get('origin', {})

    from handlers.report_base import reply_to_origin
    from database import log_report_stage
    user_id = origin.get('user_id')

    now = time.time()
    expired = [k for k, v in list(AWARD_MEDIA_CACHE.items())
               if now - v.get('created', 0) > 300]
    for k in expired:
        AWARD_MEDIA_CACHE.pop(k, None)

    if key and key in AWARD_PENDING_REPORTS and not AWARD_PENDING_REPORTS[key].get('saved'):
        pending = AWARD_PENDING_REPORTS.pop(key)
        pending['saved'] = True
        _delete_pending_from_db(key)
        merged = (pending['photos'] + photos)[:MAX_PHOTOS]
        ignored = max(0, len(pending['photos']) + len(photos) - MAX_PHOTOS)
        if ignored > 0:
            await reply_to_origin(
                context.bot, origin,
                f"⚠️ 사진 {ignored}장 무시됨 (한도 {MAX_PHOTOS}장)"
            )
        await _finalize_award(context, pending['data'], merged,
                              origin=pending.get('origin', origin))
        return

    # 앨범에서 캡션 키워드 부분 일치 안내
    if caption:
        from handlers.help_handler import is_partial_match, maybe_send_format_help
        if is_partial_match('award', caption):
            await maybe_send_format_help(context.bot, origin, user_id, 'award')

    if caption and '수상' in caption and '보고' in caption:
        log_report_stage('award', 'received', 'ok', user_id=user_id,
                         chat_id=origin.get('chat_id'),
                         topic_id=origin.get('message_thread_id'),
                         message_id=origin.get('message_id'),
                         detail=f"album {len(photos)}장")
        data = parse_award_caption(caption)
        if not data:
            log_report_stage('award', 'parsed', 'fail', user_id=user_id)
            await reply_to_origin(
                context.bot, origin,
                "⚠️ 수상보고서 파싱 실패"
            )
            return
        if '_missing' in data:
            log_report_stage('award', 'parsed', 'fail', user_id=user_id,
                             detail=f"missing: {','.join(data['_missing'])}")
            hints = '\n  - '.join(_alias_hint(f) for f in data['_missing'])
            await reply_to_origin(
                context.bot, origin,
                f"⚠️ 필수 항목 누락:\n  - {hints}"
            )
            return
        log_report_stage('award', 'parsed', 'ok', user_id=user_id)
        data['_origin'] = origin
        await reply_to_origin(
            context.bot, origin,
            f"✅ 수상보고서 접수 (사진 {len(photos)}장) - 처리 중"
        )
        await _finalize_award(context, data, photos, origin=origin)
        return

    # 캡션 없음 → PENDING_PHOTOS에 누적 보관
    if key:
        _cleanup_award_photos()
        from handlers.report_base import add_photos_to_pending, format_photo_count_msg
        added, total, ignored = add_photos_to_pending(
            AWARD_PENDING_PHOTOS, key, photos, MAX_PHOTOS,
            init_extra={'created': time.time()}
        )
        if added > 0:
            _save_pending_photos_to_db(key, AWARD_PENDING_PHOTOS[key]['photos'])
        await reply_to_origin(context.bot, origin, format_photo_count_msg(total, ignored))


async def handle_award_photo(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """사진 수신 처리 — 캡션 있음/없음, 단일/앨범 모두 대응 (케이스 A~G)"""
    if not update.message or not update.message.photo:
        return
    if update.effective_chat.id != AWARD_GROUP_ID:
        return
    if update.message.message_thread_id != AWARD_TOPIC_ID:
        return

    from handlers.report_base import make_origin, reply_to_origin
    from database import log_report_stage
    user_id = update.effective_user.id
    origin = make_origin(update)

    caption = update.message.caption or ''
    media_group_id = update.message.media_group_id
    photo = update.message.photo[-1]
    photo_file = await context.bot.get_file(photo.file_id)
    photo_url = photo_file.file_path
    key = _award_key(update.effective_chat.id,
                     update.message.message_thread_id, user_id)

    # 앨범(여러 장) 처리
    if media_group_id:
        if media_group_id not in AWARD_MEDIA_CACHE:
            AWARD_MEDIA_CACHE[media_group_id] = {
                'caption': '', 'photos': [], 'key': key, 'origin': origin,
                'processed': False, 'created': time.time()
            }
            asyncio.create_task(_process_award_album(context, media_group_id))
        AWARD_MEDIA_CACHE[media_group_id]['photos'].append(photo_url)
        if caption:
            AWARD_MEDIA_CACHE[media_group_id]['caption'] = caption
        return

    # 형식 안내: 사진+캡션이고 키워드 부분 일치 시 (1일 1회)
    if caption:
        from handlers.help_handler import is_partial_match, maybe_send_format_help
        if is_partial_match('award', caption):
            await maybe_send_format_help(context.bot, origin, user_id, 'award')

    # 케이스 A/B: 단일 사진+캡션 → 즉시 처리
    if caption and '수상' in caption and '보고' in caption:
        log_report_stage('award', 'received', 'ok', user_id=user_id,
                         chat_id=origin['chat_id'],
                         topic_id=origin['message_thread_id'],
                         message_id=origin['message_id'],
                         detail=caption[:120])
        data = parse_award_caption(caption)
        if not data:
            log_report_stage('award', 'parsed', 'fail', user_id=user_id)
            await reply_to_origin(
                context.bot, origin,
                "⚠️ 수상보고서 파싱 실패 (수상+보고 키워드 필요)"
            )
            return
        if '_missing' in data:
            log_report_stage('award', 'parsed', 'fail', user_id=user_id,
                             detail=f"missing: {','.join(data['_missing'])}")
            hints = '\n  - '.join(_alias_hint(f) for f in data['_missing'])
            await reply_to_origin(
                context.bot, origin,
                f"⚠️ 필수 항목 누락:\n  - {hints}"
            )
            return
        log_report_stage('award', 'parsed', 'ok', user_id=user_id)
        data['_origin'] = origin
        await reply_to_origin(context.bot, origin, "✅ 수상보고서 접수 - 처리 중")
        await _finalize_award(context, data, [photo_url], origin=origin)
        return

    from handlers.report_base import add_photos_to_pending, format_photo_count_msg

    # 케이스 C/D/G: 텍스트 등록 후 사진 도착
    if key in AWARD_PENDING_REPORTS and not AWARD_PENDING_REPORTS[key].get('saved'):
        added, total, ignored = add_photos_to_pending(
            AWARD_PENDING_REPORTS, key, [photo_url], MAX_PHOTOS
        )
        if added > 0:
            AWARD_PENDING_REPORTS[key]['last_photo_time'] = time.time()
            _save_pending_to_db(key, AWARD_PENDING_REPORTS[key])
        await reply_to_origin(context.bot, origin, format_photo_count_msg(total, ignored))
        return

    # 케이스 F: 사진 먼저 → PENDING_PHOTOS 보관
    _cleanup_award_photos()
    added, total, ignored = add_photos_to_pending(
        AWARD_PENDING_PHOTOS, key, [photo_url], MAX_PHOTOS,
        init_extra={'created': time.time()}
    )
    if added > 0:
        _save_pending_photos_to_db(key, AWARD_PENDING_PHOTOS[key]['photos'])
    await reply_to_origin(context.bot, origin, format_photo_count_msg(total, ignored))


async def handle_award_text(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """텍스트 보고서 수신 → PENDING 등록 (케이스 C~G 지원)"""
    if not update.message or not update.message.text:
        return
    if update.effective_chat.id != AWARD_GROUP_ID:
        return
    if update.message.message_thread_id != AWARD_TOPIC_ID:
        return
    text = update.message.text
    if '수상' not in text or '보고' not in text:
        # 부분 일치면 형식 안내 (1일 1회)
        from handlers.help_handler import is_partial_match, maybe_send_format_help
        from handlers.report_base import make_origin as _mo
        if is_partial_match('award', text):
            await maybe_send_format_help(context.bot, _mo(update),
                                          update.effective_user.id, 'award')
        return

    from handlers.report_base import make_origin, reply_to_origin
    from database import log_report_stage
    user_id = update.effective_user.id
    origin = make_origin(update)

    log_report_stage('award', 'received', 'ok',
                     user_id=user_id, chat_id=origin.get('chat_id'),
                     topic_id=origin.get('message_thread_id'),
                     message_id=origin.get('message_id'),
                     detail=text[:120])

    data = parse_award_caption(text)
    if not data:
        log_report_stage('award', 'parsed', 'fail',
                         user_id=user_id, detail='parse returned None')
        await reply_to_origin(
            context.bot, origin,
            "⚠️ 수상보고서로 인식하지 못했습니다.\n"
            "첫 줄 형식: '지역 지부 수상보고' / 키워드 '수상'+'보고' 필요"
        )
        return
    if '_missing' in data:
        log_report_stage('award', 'parsed', 'fail',
                         user_id=user_id,
                         detail=f"missing: {','.join(data['_missing'])}")
        hints = '\n  - '.join(_alias_hint(f) for f in data['_missing'])
        await reply_to_origin(
            context.bot, origin,
            f"⚠️ 필수 항목 누락:\n  - {hints}"
        )
        return

    log_report_stage('award', 'parsed', 'ok', user_id=user_id,
                     detail=f"{data.get('지부','')} | {data.get('수상명','')}")

    key = _award_key(update.effective_chat.id,
                     update.message.message_thread_id, user_id)
    # 케이스 F/G: 먼저 도착한 사진 연결
    pre_photos: list = []
    if key in AWARD_PENDING_PHOTOS:
        pre_photos = AWARD_PENDING_PHOTOS.pop(key)['photos']
        try:
            from database import delete_pending_photos_db
            delete_pending_photos_db('award', _key_to_str(key))
        except Exception:
            pass

    data['_origin'] = origin   # 영속화 시 같이 저장됨

    AWARD_PENDING_REPORTS[key] = {
        'data': data, 'photos': pre_photos, 'saved': False,
        'last_photo_time': time.time() if pre_photos else 0,
        'created': time.time(),
        'origin': origin,
    }
    _save_pending_to_db(key, AWARD_PENDING_REPORTS[key])

    # 보고자 즉시 ack
    photo_msg = f" (이미 받은 사진 {len(pre_photos)}장 포함)" if pre_photos else ""
    await reply_to_origin(
        context.bot, origin,
        f"✅ 수상보고서 접수{photo_msg}\n사진 대기 중 (60s+ 추가 5분 자동 연장)"
    )

    asyncio.create_task(_flush_award_report(context, key))

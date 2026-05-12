"""3개 보고서(봉사/MOU/수상) 공통 베이스 유틸리티"""

import asyncio
import tempfile
import time

import pytz
import requests
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

import config

KST = pytz.timezone('Asia/Seoul')
ADMIN_USER_ID = 97057565


# ── 텔레그램 통신 ──────────────────────────────────────────────────────────────

async def notify_admin(bot, message: str):
    """관리자 백업 알림 (서무 DM 실패 시 등)"""
    try:
        await bot.send_message(chat_id=ADMIN_USER_ID, text=message[:4000])
    except Exception as e:
        print(f"❌ 관리자 알림 실패: {e}")


async def send_to_recipient(bot, recipient_id: int, retries: int = 3, **kwargs):
    """서무 DM 전송. 실패 시 1초·2초 백오프 후 재시도. 최종 실패 시 관리자 백업."""
    last_err = None
    for attempt in range(retries):
        try:
            if 'document' in kwargs:
                await bot.send_document(chat_id=recipient_id, **kwargs)
            else:
                await bot.send_message(chat_id=recipient_id, **kwargs)
            return True
        except Exception as e:
            last_err = e
            if attempt < retries - 1:
                await asyncio.sleep(2 ** attempt)  # 1, 2, 4초
    await notify_admin(
        bot,
        f"❌ 서무({recipient_id}) DM {retries}회 실패: {last_err}\n"
        f"내용: {str(kwargs.get('text', ''))[:300]}"
    )
    return False


async def reply_to_origin(bot, origin: dict, text: str) -> bool:
    """보고자 원본 메시지에 답글. origin = {chat_id, message_id, message_thread_id}"""
    if not origin or not origin.get('chat_id') or not origin.get('message_id'):
        return False
    try:
        kwargs = {
            'chat_id': origin['chat_id'],
            'text': text[:4000],
            'reply_to_message_id': origin['message_id'],
        }
        if origin.get('message_thread_id'):
            kwargs['message_thread_id'] = origin['message_thread_id']
        await bot.send_message(**kwargs)
        return True
    except Exception as e:
        print(f"⚠️ 보고자 답글 실패: {e}")
        return False


def make_origin(update) -> dict:
    """telegram Update에서 보고자 메시지 위치 추출"""
    if not update or not update.message:
        return {}
    return {
        'chat_id': update.effective_chat.id,
        'message_id': update.message.message_id,
        'message_thread_id': update.message.message_thread_id,
        'user_id': update.effective_user.id if update.effective_user else None,
    }


def with_sheet_retry(save_fn, data, retries: int = 3) -> bool:
    """시트 저장 재시도. save_fn(data) 호출 — 항상 1 인자.
    실패 시 1·2·4초 지수 백오프."""
    for attempt in range(retries):
        try:
            if save_fn(data):
                return True
        except Exception as e:
            print(f"⚠️ 시트 저장 attempt {attempt+1} 실패: {e}")
        if attempt < retries - 1:
            time.sleep(2 ** attempt)
    return False


# ── 사진 다운로드 ──────────────────────────────────────────────────────────────

def download_photo(url: str, retries: int = 3, delay: int = 5) -> str | None:
    """사진 다운로드 (최대 retries회, 실패 시 delay초 대기 후 재시도)"""
    token = config.get('telegram_token')
    full_url = url if url.startswith('http') else f"https://api.telegram.org/file/bot{token}/{url}"
    last_err = None
    for attempt in range(retries):
        try:
            resp = requests.get(full_url, timeout=10)
            if resp.status_code == 200:
                suffix = '.jpg' if 'jpg' in url.lower() else '.png'
                tmp = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
                tmp.write(resp.content)
                tmp.close()
                return tmp.name
            last_err = f"HTTP {resp.status_code}"
        except Exception as e:
            last_err = str(e)
        if attempt < retries - 1:
            time.sleep(delay)
    print(f"❌ 사진 다운로드 {retries}회 모두 실패 ({last_err}): {url[:60]}")
    return None


async def download_photos_batch(photo_urls: list) -> tuple:
    """사진 다운로드. 반환: (성공 임시파일 리스트, 실패 개수)"""
    loop = asyncio.get_running_loop()
    tmp_files = []
    failed = 0
    for url in photo_urls:
        if not url:
            continue
        tmp = await loop.run_in_executor(None, download_photo, url)
        if tmp:
            tmp_files.append(tmp)
        else:
            failed += 1
    return tmp_files, failed


# ── Word docx 헬퍼 ─────────────────────────────────────────────────────────────

def set_cell_bg(cell, color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)


def add_photos_grid(doc, photo_paths: list, title: str = '사진'):
    """사진 2열 그리드 (8cm x 5cm). 1페이지에 4장."""
    if not photo_paths:
        return
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(8)
    for i in range(0, len(photo_paths), 2):
        pair = photo_paths[i:i+2]
        photo_table = doc.add_table(rows=1, cols=len(pair))
        for j, path in enumerate(pair):
            cell = photo_table.rows[0].cells[j]
            try:
                para = cell.paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                pic_run = para.add_run()
                pic_run.add_picture(path, width=Cm(8), height=Cm(5))
            except Exception as e:
                print(f"❌ 사진 삽입 오류: {e}")
                cell.paragraphs[0].add_run("사진 오류")
        doc.add_paragraph()


# ── PENDING 대기 관리 ──────────────────────────────────────────────────────────

def cleanup_pending_photos(pending_dict: dict, ttl: int = 300):
    now = time.time()
    expired = [k for k, v in list(pending_dict.items())
               if now - v.get('created', 0) > ttl]
    for k in expired:
        pending_dict.pop(k, None)


# 보고서당 최대 사진 수
MAX_PHOTOS = 10


def add_photos_to_pending(pending_dict: dict, key, new_photos: list,
                          max_total: int = MAX_PHOTOS,
                          init_extra: dict = None) -> tuple:
    """
    pending_dict[key]['photos']에 사진 누적, max_total 도달 시 무시.
    pending_dict[key] 없으면 생성 (init_extra로 추가 필드 설정 가능).
    Returns: (added_count, total_after, ignored_count)
    """
    if key not in pending_dict:
        pending_dict[key] = {'photos': []}
        if init_extra:
            pending_dict[key].update(init_extra)
    current_photos = pending_dict[key]['photos']
    if len(current_photos) >= max_total:
        return 0, len(current_photos), len(new_photos)
    can_add = max_total - len(current_photos)
    to_add = new_photos[:can_add]
    current_photos.extend(to_add)
    return len(to_add), len(current_photos), len(new_photos) - len(to_add)


def format_photo_count_msg(total: int, ignored: int = 0,
                            max_total: int = MAX_PHOTOS) -> str:
    """사진 누적 카운트 메시지"""
    if total >= max_total and ignored > 0:
        return f"📸 {total}/{max_total}장 접수 (추가 {ignored}장 무시됨)"
    if total >= max_total:
        return f"📸 {total}/{max_total}장 접수 (이후 추가 사진은 무시됩니다)"
    return f"📸 {total}/{max_total}장 접수"


# ── 중복 제출 감지 (dedup) ────────────────────────────────────────────────────

def normalize_for_hash(s: str) -> str:
    """hash 생성용 정규화: 공백 제거 + 소문자.
    '대전 지부' vs '대전지부' 같은 표기 차이를 흡수."""
    if not s:
        return ''
    return ''.join(s.split()).lower()


def build_submission_hash(report_type: str, data: dict) -> str:
    """보고서별 강화된 dedup hash"""
    if report_type == 'award':
        parts = [data.get('지부', ''), data.get('수상명', ''), data.get('수상자', '')]
    elif report_type == 'mou':
        parts = [data.get('지부', ''), data.get('협약명', ''), data.get('기관명', '')]
    elif report_type == 'service':
        parts = [
            data.get('지파명', ''), data.get('교회명', ''),
            data.get('활동명', ''),
            (data.get('활동일시', '') or '')[:10],
        ]
    else:
        parts = []
    return '|'.join(normalize_for_hash(p) for p in parts)


_REPORT_LABEL = {'award': '수상', 'mou': 'MOU', 'service': '봉사'}


async def check_duplicate_and_warn(context, *, report_type: str, data: dict,
                                    origin: dict, recipient_id: int) -> tuple:
    """dedup 체크 (옵션 3: 강제 진행 + 경고).
    Returns: (sub_hash, was_duplicate)"""
    user_id = origin.get('user_id') if origin else None
    sub_hash = build_submission_hash(report_type, data)
    if not sub_hash or sub_hash.replace('|', '') == '':
        return sub_hash, False

    from database import find_recent_submission
    existing = find_recent_submission(report_type, sub_hash, window_sec=600)
    if not existing:
        return sub_hash, False

    same_user = bool(user_id) and existing.get('user_id') == user_id
    minutes_ago = max(1, int((time.time() - existing.get('submitted_at', 0)) / 60))
    label = _REPORT_LABEL.get(report_type, report_type)
    prev_summary = (existing.get('summary') or '')[:150]

    # 보고자 안내 (강/약 경고 차등)
    if same_user:
        boggoja_msg = (
            f"⚠️ 중복 의심 — {minutes_ago}분 전 동일 보고서 처리됨\n"
            f"📌 같은 분이 동일 내용 제출 ({label} 보고서)\n\n"
            f"✅ 새 보고서가 맞으면 그대로 두세요 (정상 처리됨)\n"
            f"❌ 잘못 보냈으면 서무에게 시트 삭제 요청"
        )
    else:
        boggoja_msg = (
            f"⚠️ 중복 의심 — {minutes_ago}분 전 다른 분이 동일 보고서 제출\n"
            f"📌 같은 행사를 두 명이 보고했을 수 있음 ({label} 보고서)\n\n"
            f"✅ 새 보고서가 맞으면 그대로 두세요"
        )
    await reply_to_origin(context.bot, origin, boggoja_msg)

    # 서무 안내
    if same_user:
        seomu_msg = (
            f"⚠️ 중복 의심 보고서 — 같은 사용자\n"
            f"종류: {label}\n"
            f"이전 ({minutes_ago}분 전): {prev_summary}\n"
            f"💡 시트 확인 후 중복이면 삭제 권장"
        )
    else:
        seomu_msg = (
            f"⚠️ 중복 의심 보고서 — 다른 사용자\n"
            f"종류: {label}\n"
            f"이전 ({minutes_ago}분 전): {prev_summary}\n"
            f"💡 같은 행사 두 사람 보고 가능성 (중복 시 시트 정리)"
        )
    await send_to_recipient(context.bot, recipient_id, text=seomu_msg)

    return sub_hash, True

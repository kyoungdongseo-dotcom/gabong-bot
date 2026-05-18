import os
import requests
import tempfile
import pytz
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import config

KST = pytz.timezone('Asia/Seoul')
DOCX_RECIPIENT_ID = config.get('secretary_id', 754270008)

def set_cell_background(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)

def add_table_row(table, label, value):
    row = table.add_row()
    label_cell = row.cells[0]
    value_cell = row.cells[1]
    label_cell.text = label
    label_cell.paragraphs[0].runs[0].bold = True
    label_cell.paragraphs[0].runs[0].font.size = Pt(10)
    set_cell_background(label_cell, 'D5E8F0')
    value_cell.text = value or '-'
    value_cell.paragraphs[0].runs[0].font.size = Pt(10)

def add_section(doc, title, content):
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    content_p = doc.add_paragraph(content or '-')
    content_p.runs[0].font.size = Pt(10)
    content_p.paragraph_format.space_after = Pt(8)

def download_photo(url: str) -> str | None:
    try:
        token = config.get_telegram_token()
        if not url.startswith('http'):
            full_url = f"https://api.telegram.org/file/bot{token}/{url}"
        else:
            full_url = url
        response = requests.get(full_url, timeout=10)
        if response.status_code == 200:
            suffix = '.jpg' if 'jpg' in url.lower() else '.png'
            tmp = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
            tmp.write(response.content)
            tmp.close()
            return tmp.name
        return None
    except Exception as e:
        print(f"❌ 사진 다운로드 오류: {e}")
        return None

def add_photos_grid(doc, photo_paths):
    """사진을 2열로 배치 (1페이지 4장)"""
    if not photo_paths:
        return

    p = doc.add_paragraph()
    run = p.add_run('📸 봉사 사진')
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(8)

    # 2장씩 같은 줄에 배치
    for i in range(0, len(photo_paths), 2):
        pair = photo_paths[i:i+2]

        # 2열 테이블 생성
        photo_table = doc.add_table(rows=1, cols=len(pair))

        for j, path in enumerate(pair):
            cell = photo_table.rows[0].cells[j]
            try:
                para = cell.paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run()
                run.add_picture(path, width=Cm(8), height=Cm(5))
            except Exception as e:
                print(f"❌ 사진 삽입 오류: {e}")
                cell.paragraphs[0].add_run("사진 오류")

        doc.add_paragraph()

def generate_docx(report: dict, output_path: str, photo_paths: list = None) -> bool:
    """photo_paths: 호출 측이 미리 다운로드한 임시 파일 경로 리스트.
    None 이면 report 의 사진URL 에서 직접 다운로드 (하위 호환)."""
    try:
        doc = Document()

        for section in doc.sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 신양식: "{연합회}연합회 {지부} 봉사 활동보고"
        # 구양식: "{지파명} {교회명} 봉사보고서" (연합회/지부 정규화 실패 시 fallback)
        if report.get('_format') == 'new' and report.get('연합회') and report.get('지부'):
            title_text = f"{report.get('연합회')}연합회 {report.get('지부')} 봉사 활동보고"
        else:
            title_text = f"{report.get('지파명', '')} {report.get('교회명', '')} 봉사보고서"
        run = title.add_run(title_text)
        run.bold = True
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        title.paragraph_format.space_after = Pt(6)

        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = subtitle.add_run(report.get('활동명', ''))
        run2.font.size = Pt(12)
        run2.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
        subtitle.paragraph_format.space_after = Pt(12)

        table = doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'
        table.columns[0].width = Cm(4)
        table.columns[1].width = Cm(12)

        add_table_row(table, '■ 활동명', report.get('활동명'))
        add_table_row(table, '■ 봉사분류', report.get('봉사분류'))
        add_table_row(table, '■ 활동일시', report.get('활동일시'))
        add_table_row(table, '■ 활동장소', report.get('활동장소'))
        add_table_row(table, '■ 수혜자 수', report.get('수혜자수'))
        if report.get('수혜기관대상'):
            add_table_row(table, '■ 수혜기관/대상', report.get('수혜기관대상'))
        add_table_row(table, '■ 내부봉사자', f"{report.get('내부봉사자', '0')}명")
        add_table_row(table, '■ 외부봉사자', f"{report.get('외부봉사자', '0')}명")
        add_table_row(table, '■ 총봉사자', f"{report.get('총봉사자', '0')}명")
        if report.get('캠페인시민참여'):
            add_table_row(table, '■ 캠페인 시민참여', report.get('캠페인시민참여'))
        if report.get('쓰레기수거량'):
            add_table_row(table, '■ 쓰레기 수거량', report.get('쓰레기수거량'))

        doc.add_paragraph()

        # 자유 서술 섹션 — 값이 있는 항목만 출력 (신/구 양식 양방향 호환, 2026-05-18)
        idx = 1
        if report.get('기획취지배경'):
            add_section(doc, f'{idx}. 기획 취지·배경(지역문제점)', report.get('기획취지배경'))
            idx += 1
        if report.get('활동내용'):
            add_section(doc, f'{idx}. 활동 내용', report.get('활동내용'))
            idx += 1
        if report.get('활동성과'):
            add_section(doc, f'{idx}. 활동 성과', report.get('활동성과'))
            idx += 1
        # 협력 인사/단체 (수치 + 소속) 통합 섹션
        cooperation_parts = []
        if report.get('협력인사수'):
            cooperation_parts.append(f"인사 {report.get('협력인사수')}명")
        if report.get('협력단체수'):
            cooperation_parts.append(f"단체 {report.get('협력단체수')}곳")
        if report.get('협력소속'):
            cooperation_parts.append(f"소속: {report.get('협력소속')}")
        if cooperation_parts:
            add_section(doc, f'{idx}. 협력 인사/단체', ' | '.join(cooperation_parts))
            idx += 1
        # 현장 반응 (수혜자 / 시민 참여)
        reactions = []
        if report.get('수혜자반응'):
            reactions.append(f"수혜자 반응: {report.get('수혜자반응')}")
        if report.get('시민참여반응'):
            reactions.append(f"시민 참여 반응: {report.get('시민참여반응')}")
        if reactions:
            add_section(doc, f'{idx}. 현장 반응', '\n'.join(reactions))
            idx += 1
        if report.get('기대효과'):
            add_section(doc, f'{idx}. 기대효과(지역사회관점)', report.get('기대효과'))
            idx += 1
        # 구양식 잔존 필드 — 호환 모드에서만 출력
        if report.get('반응특이사항'):
            add_section(doc, f'{idx}. 반응 및 특이사항', report.get('반응특이사항'))
            idx += 1
        if report.get('참여인사'):
            add_section(doc, f'{idx}. 참여인사', report.get('참여인사'))
            idx += 1
        if report.get('홍보도구'):
            add_section(doc, f'{idx}. 홍보도구', report.get('홍보도구'))
            idx += 1
        if report.get('잘된점'):
            add_section(doc, f'{idx}. 잘된 점', report.get('잘된점'))
            idx += 1
        if report.get('개선할점'):
            add_section(doc, f'{idx}. 개선할 점', report.get('개선할점'))
            idx += 1

        # 사진 처리: photo_paths 가 있으면 그대로, 없으면 URL 에서 다운로드
        if photo_paths is not None:
            valid_paths = [p for p in photo_paths if p and os.path.exists(p)]
            if valid_paths:
                add_photos_grid(doc, valid_paths)
            # 임시 파일 삭제는 호출 측 책임
        else:
            photo_urls = [report.get(f'사진{i}링크', '') for i in range(1, 11)
                          if report.get(f'사진{i}링크')]
            if photo_urls:
                tmp_files = []
                failed_count = 0
                for url in photo_urls:
                    tmp_path = download_photo(url)
                    if tmp_path:
                        tmp_files.append(tmp_path)
                    else:
                        failed_count += 1
                if tmp_files:
                    add_photos_grid(doc, tmp_files)
                if failed_count > 0:
                    note_p = doc.add_paragraph()
                    note_run = note_p.add_run(
                        f"⚠️ 사진 {len(tmp_files)}장 첨부 완료 / {failed_count}장 다운로드 실패"
                        " (텔레그램 파일 링크 만료 가능)"
                    )
                    note_run.font.size = Pt(9)
                    note_run.font.color.rgb = RGBColor(0xFF, 0x66, 0x00)
                for tmp_path in tmp_files:
                    try:
                        os.remove(tmp_path)
                    except Exception:
                        pass

        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run3 = footer.add_run(f"등록일시: {report.get('등록일시', '')}")
        run3.font.size = Pt(9)
        run3.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        doc.save(output_path)
        print(f"✅ Word 파일 생성 완료: {output_path}")
        return True

    except Exception as e:
        print(f"❌ Word 생성 오류: {e}")
        return False

# generate_and_send_docx 는 트랜잭션 패턴 도입 후 사용처 0건 → 제거.
# 봉사보고서 finalize_report 가 generate_docx 직접 호출 + 사진 paths 전달.
